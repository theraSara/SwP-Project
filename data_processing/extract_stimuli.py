import re
import math
import os
import pandas as pd
from docx import Document

N_PARTICIPANTS = 216
S_FACTOR = 200  # Nair & Oh smoothing
RT_SOURCE = "data/SPRT_LogLin_216.csv"
DOCX_SOURCE = "data/Stimuli_Appendix_format.docx"
OUTPUT_FILE = "data_output/bk21_stimuli_final.csv"

# Load RT file for item mapping
rt = pd.read_csv(RT_SOURCE)
id_map = {}
for _, row in rt.iterrows():
    cond_key = row['condition'][0].upper()
    id_map[(row['critical_word'].lower(), cond_key)] = row['ITEM']

doc = Document(DOCX_SOURCE)
table = doc.tables[0]
rows = []

for row in table.rows:
    cell = row.cells[0]
    raw_text = cell.text.strip().replace("*", "")

    if not raw_text.startswith(("H:", "M:", "L:")):
        continue

    # Extract cloze %
    cloze_match = re.search(r"\((\d+(?:\.\d+)?)%?\)\s*$", raw_text)
    cloze_percent = float(cloze_match.group(1)) if cloze_match else 0.0

    # Extract underlined critical word
    critical_word_raw = ""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.underline:
                critical_word_raw += run.text

    critical_word = critical_word_raw.strip().strip(".,;!?'\"")
    cond_letter = raw_text[0].upper()

    # Strip leading label and trailing cloze
    sentence_full = re.sub(r"\((\d+(?:\.\d+)?)%?\)\s*$", "", raw_text[2:]).strip()

    # Find target span
    match = re.search(r'\b' + re.escape(critical_word) + r'\b', sentence_full)
    if match:
        start, end = match.span()
        prefix = sentence_full[:start].rstrip()
        suffix = sentence_full[end:].lstrip()
    else:
        parts = sentence_full.split(critical_word, 1)
        prefix = parts[0].rstrip()
        suffix = parts[1].lstrip() if len(parts) > 1 else ""

    # Cloze-derived measures
    cloze_prob = cloze_percent / 100.0
    count = cloze_prob * N_PARTICIPANTS

    # Nair & Oh-style smoothing
    cloze_prob_nair_oh = (count + 1) / (N_PARTICIPANTS + S_FACTOR)
    cloze_surprisal_bits_nair_oh = -math.log2(cloze_prob_nair_oh)
    cloze_log10_nair_oh = math.log10(cloze_prob_nair_oh)

    # Conventional add-one smoothing
    cloze_prob_add1 = (count + 1) / (N_PARTICIPANTS + 1)
    cloze_surprisal_bits_add1 = -math.log2(cloze_prob_add1)
    cloze_log10_add1 = math.log10(cloze_prob_add1)

    true_item = id_map.get((critical_word.lower(), cond_letter), "MISSING")

    rows.append({
        "ITEM": true_item,
        "condition": {"H": "HC", "M": "MC", "L": "LC"}[cond_letter],
        "cloze_percent": cloze_percent,
        "cloze_prob": cloze_prob,
        "cloze_count": count,

        "cloze_prob_nair_oh": cloze_prob_nair_oh,
        "cloze_surprisal_bits_nair_oh": cloze_surprisal_bits_nair_oh,
        "cloze_log10_nair_oh": cloze_log10_nair_oh,

        "cloze_prob_add1": cloze_prob_add1,
        "cloze_surprisal_bits_add1": cloze_surprisal_bits_add1,
        "cloze_log10_add1": cloze_log10_add1,

        "sentence": sentence_full,
        "prefix": prefix,
        "suffix": suffix,
        "critical_word": critical_word,
        "target_llm": critical_word
    })

df = pd.DataFrame(rows)
os.makedirs("data_output", exist_ok=True)
df.to_csv(OUTPUT_FILE, index=False)

print(f"Successfully created {OUTPUT_FILE}")
print(f"Total items: {len(df)} | Missing IDs: {len(df[df['ITEM'] == 'MISSING'])}")