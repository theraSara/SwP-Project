import re
import math
import os
import pandas as pd
from docx import Document

N_PARTICIPANTS = 216
S_FACTOR = 200  # Best-fit benchmark from Nair & Oh (2026)
RT_SOURCE = "data/SPRT_LogLin_216.csv"
DOCX_SOURCE = "data/Stimuli_Appendix_format.docx"
OUTPUT_FILE = "data_output/bk21_stimuli_final.csv"

# Load the RT file to create the ID Source of Truth
rt = pd.read_csv(RT_SOURCE)
id_map = {}
for _, row in rt.iterrows():
    cond_key = row['condition'][0].upper() 
    id_map[(row['critical_word'].lower(), cond_key)] = row['ITEM']

doc = Document(DOCX_SOURCE)
table = doc.tables[0]
rows = []

for row in table.rows:
    cell = row.cells[0]
    raw_text = cell.text.strip().replace("*", "")
    if not raw_text.startswith(("H:", "M:", "L:")):
        continue

    # extract Cloze % from parentheses
    cloze_match = re.search(r"\((\d+(?:\.\d+)?)%?\)\s*$", raw_text)
    cloze_percent = float(cloze_match.group(1)) if cloze_match else 0.0

    # underlined Critical Word
    critical_word_raw = ""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.underline:
                critical_word_raw += run.text
    critical_word = critical_word_raw.strip().strip(".,;!?'\"")
    
    cond_letter = raw_text[0].upper()
    
    sentence_full = re.sub(r"\((\d+(?:\.\d+)?)%?\)\s*$", "", raw_text[2:]).strip()
    
    # word boundaries 
    match = re.search(r'\b' + re.escape(critical_word) + r'\b', sentence_full)
    if match:
        start, end = match.span()
        prefix = sentence_full[:start].rstrip()
        suffix = sentence_full[end:].lstrip()
    else:
        # Fallback
        parts = sentence_full.split(critical_word, 1) 
        prefix = parts[0].rstrip()
        suffix = parts[1].lstrip() if len(parts) > 1 else ""

    # Apply Nair & Oh (2026) Smoothing (S=200, Add-1) = -log2( (count + 1) / (N + S) )
    count = (cloze_percent / 100.0) * N_PARTICIPANTS
    cloze_surprisal = -math.log2((count + 1) / (N_PARTICIPANTS + S_FACTOR))

    true_item = id_map.get((critical_word.lower(), cond_letter), "MISSING")

    rows.append({
        "ITEM": true_item,
        "condition": {"H": "HC", "M": "MC", "L": "LC"}[cond_letter],
        "cloze_percent": cloze_percent,
        "cloze_surprisal": cloze_surprisal,
        "sentence": sentence_full,
        "prefix": prefix,
        "suffix": suffix,
        "critical_word": critical_word,
        "target_llm": critical_word
    })

df = pd.DataFrame(rows)
os.makedirs("data_output", exist_ok=True)
df.to_csv(OUTPUT_FILE, index=False)

print(f"Successfully created {OUTPUT_FILE}")
print(f"Total items: {len(df)} | Missing IDs: {len(df[df['ITEM'] == 'MISSING'])}")